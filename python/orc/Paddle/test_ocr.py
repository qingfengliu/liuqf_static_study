import os

from pdf2image import convert_from_path
from paddleocr import PaddleOCR
import numpy as np
from docx import Document

import sys
sys.path.append(r'D:\program\Release-24.08.0-0\poppler-24.08.0\Library\bin')
father_path = 'D:\工作'
input_path = os.path.join(father_path, '浙江省人民政府办公厅关于印发浙江省推动碳排双控工作若干举措的通知.pdf')
output_path = os.path.join(father_path, 'word/浙江省人民政府办公厅关于印发浙江省推动碳排双控工作若干举措的通知.docx')
# 步骤 1: 将 PDF 转换为图片
pages = convert_from_path(input_path, 300)  # 转换为 PIL 图像对象

# 步骤 2: 初始化 OCR 模型
ocr = PaddleOCR(use_angle_cls=True, lang='ch')

# 步骤 3: 遍历每一页，进行 OCR 处理
results = []
for page in pages:
    # 将 PIL 图像转换为 numpy 数组
    page_np = np.array(page)

    # 使用 OCR 提取文本
    ocr_result = ocr.ocr(page_np, cls=True)
    page_text = []
    for line in ocr_result:
        if line:  # 检查 line 是否为 None
            for word_info in line:
                # word_info 包含了文本和置信度
                text, confidence = word_info[1]
                page_text.append(text)  # 提取文本部分

    results.append('\n'.join(page_text))

# 打印提取的文本
# for result in results:
#     print(result)

# 步骤 4: 将文本保存为 DOCX 文件
doc = Document()
for i, page_text in enumerate(results, 1):
    doc.add_heading(f'Page {i}', level=1)
    doc.add_paragraph(page_text)

doc.save(output_path)
